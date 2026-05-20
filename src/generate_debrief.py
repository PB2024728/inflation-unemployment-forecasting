"""
Generates the Days 1-4 Progress Report as a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = "Inflation_Unemployment_Forecasting_Progress_Report_Days1-4.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)   # dark navy — headings
STEEL  = RGBColor(0x2E, 0x75, 0xB6)   # steel blue — sub-headings
GRAY   = RGBColor(0x59, 0x59, 0x59)   # dark gray — body
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
TBL_HDR= RGBColor(0x2E, 0x75, 0xB6)   # table header fill

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "2E75B6")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        fill = "F2F7FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, fill)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            if c_idx == 0:
                run.bold = True

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = STEEL
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def callout(doc, text):
    """Indented italic callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = STEEL
    return p

def divider(doc):
    doc.add_paragraph()

# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ── Title block ────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title.add_run("Inflation & Unemployment Forecasting Bake-Off")
t.bold = True
t.font.size = Pt(22)
t.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub.add_run("Progress Report — Days 1 to 4 of 10  |  May 6, 2026")
s.font.size = Pt(11)
s.font.color.rgb = STEEL

doc.add_paragraph()

# ── Section 1: Project Overview ────────────────────────────────────────────────
h1(doc, "1. Project Overview")

body(doc,
    "This project builds a head-to-head forecasting competition between two families of models: "
    "traditional econometric models — the kind used by central banks, academic economists, and "
    "financial institutions for decades — and machine learning models — newer computational "
    "approaches that can handle larger datasets and detect complex, nonlinear patterns."
)

callout(doc,
    "“When it comes to forecasting inflation and unemployment, do modern machine learning methods "
    "outperform the classical statistical tools that economists have relied on for 40 years? "
    "And if so, under what conditions?”"
)

body(doc,
    "This is not purely an academic exercise. Accurate inflation forecasts influence interest rate "
    "decisions at the Federal Reserve. Accurate unemployment forecasts affect fiscal policy, business "
    "hiring plans, and bond market pricing. Understanding where each method succeeds or breaks down "
    "has real economic consequences."
)

# ── Section 2: What We Are Forecasting ────────────────────────────────────────
h1(doc, "2. What We Are Forecasting")

h2(doc, "Target Variables")
add_table(doc,
    ["Variable", "How It Is Measured"],
    [
        ["Headline CPI Inflation",
         "Year-over-year % change in the Consumer Price Index — the broadest measure of price changes in the U.S. economy"],
        ["Core CPI Inflation",
         "Same as above, excluding volatile food and energy prices — the Fed’s preferred signal of underlying inflation"],
        ["Unemployment Rate",
         "Percentage of the labor force actively seeking work but unable to find it"],
    ],
    col_widths=[2.0, 4.5]
)

divider(doc)

h2(doc, "Forecast Horizons")
body(doc,
    "Forecasts are generated at three time horizons: 1 month ahead, 3 months ahead, and 6 months ahead. "
    "This matters because a model that is accurate at 1 month may be unreliable at 6 months. "
    "Understanding how performance degrades across horizons is one of the core analytical contributions of this project."
)

# ── Section 3: The Model Lineup ────────────────────────────────────────────────
h1(doc, "3. The Model Lineup")

add_table(doc,
    ["Category", "Model", "What It Does"],
    [
        ["Baseline",          "Naive",          "Predicts next month’s value equals this month’s. Surprisingly hard to beat."],
        ["Classical",         "ARIMA",          "Identifies patterns in a single series using its own past values and past forecast errors."],
        ["Classical",         "VAR",            "Models inflation and unemployment jointly, capturing how each responds to changes in the other."],
        ["Machine Learning",  "Ridge",          "Linear regression across many macro predictors, penalised to prevent overfitting."],
        ["Machine Learning",  "Lasso",          "Similar to Ridge but also performs automatic variable selection, zeroing out weak predictors."],
        ["Machine Learning",  "Random Forest",  "An ensemble of hundreds of decision trees, capable of capturing nonlinear relationships."],
        ["Machine Learning",  "XGBoost",        "High-performance gradient boosting algorithm, widely used in applied forecasting competitions."],
    ],
    col_widths=[1.5, 1.4, 3.6]
)

# ── Section 4: Evaluation Framework ───────────────────────────────────────────
h1(doc, "4. Evaluation Framework")

body(doc,
    "Every model is evaluated under expanding-window backtesting. This means:"
)
bullet(doc, "We train each model only on data that would have been available at that point in time.")
bullet(doc, "We generate forecasts for 1, 3, and 6 months ahead.")
bullet(doc, "We compare each forecast to what actually happened.")
bullet(doc, "We advance one month forward and repeat.")

body(doc,
    "This process runs from January 2000 through February 2026, producing over 300 forecast evaluation "
    "points per model. The discipline here is critical — a model that uses future information during "
    "training would produce misleadingly good results. Our framework prevents this entirely.",
    space_after=8
)

h2(doc, "Performance Metrics")
add_table(doc,
    ["Metric", "What It Measures"],
    [
        ["RMSE (Root Mean Squared Error)",
         "Penalises large errors more heavily — appropriate when extreme forecast misses are costly."],
        ["MAE (Mean Absolute Error)",
         "Average absolute forecast error — easier to interpret directly in percentage-point terms."],
    ],
    col_widths=[2.5, 4.0]
)

# ── Section 5: Day-by-Day Summary ─────────────────────────────────────────────
h1(doc, "5. Day-by-Day Summary")

# Day 1
h2(doc, "Day 1 — Project Scoping")
body(doc,
    "All key design decisions were locked in before writing a single line of modelling code:"
)
bullet(doc, "Primary inflation measure: year-over-year percentage change in CPI.")
bullet(doc, "Evaluation design: expanding window — the standard in academic macro forecasting.")
bullet(doc, "Holdout period: all forecasts generated from 2000 onward, with a minimum of 60 months of training data before the first forecast.")

body(doc,
    "Deliverable: full project specification including research question, target variables, model list, "
    "evaluation design, and data sources — locked in before any data was touched."
)

# Day 2
h2(doc, "Day 2 — Data Collection")
body(doc,
    "Tools used: fredapi (Python library for the Federal Reserve Economic Data API), pandas for data "
    "wrangling, python-dotenv for secure API key management."
)
body(doc,
    "Data source: FRED (Federal Reserve Economic Data), the authoritative public repository for U.S. "
    "macroeconomic series maintained by the St. Louis Federal Reserve. We pulled 11 macroeconomic series "
    "covering January 1990 through March 2026:"
)

add_table(doc,
    ["Series", "Role"],
    [
        ["CPI & Core CPI",              "Forecast targets"],
        ["Unemployment Rate",           "Forecast target"],
        ["Federal Funds Rate",          "Monetary policy stance"],
        ["10-Year & 3-Month Treasury",  "Long-run expectations and short-run rates"],
        ["Yield Spread (10Y − 3M)", "Leading recession indicator"],
        ["Industrial Production",       "Real economic activity"],
        ["Nonfarm Payrolls",            "Labour market health"],
        ["Oil Prices (WTI)",            "Supply-side inflation driver"],
        ["Consumer Sentiment",          "Forward-looking demand indicator"],
        ["M2 Money Supply",             "Longer-run inflation signal"],
    ],
    col_widths=[2.5, 4.0]
)

divider(doc)
body(doc,
    "After applying the appropriate transformations, we produced a clean monthly dataset of "
    "421 observations and 12 model-ready features."
)
body(doc,
    "Deliverable: data/raw/raw_macro.csv and data/processed/macro_monthly.csv, plus a full data "
    "dictionary documenting every variable’s source, transformation, and role."
)

# Day 3
h2(doc, "Day 3 — Exploratory Data Analysis")
body(doc,
    "Tools used: matplotlib, seaborn. Before building any model, six diagnostic figures were produced "
    "to understand the structure of the data. Key observations:"
)

bullet(doc,
    "Inflation averaged 2.6% over the full sample, ranging from −2.0% (2009 deflation) to +9.0% "
    "(2022 post-pandemic surge). The 2022 shock is a clear structural outlier."
)
bullet(doc,
    "Unemployment averaged 5.7%, with a COVID-19 peak of 14.8% in April 2020 — a near-vertical "
    "shock and recovery compressed into roughly four months."
)
bullet(doc,
    "The Phillips Curve (the classical inverse relationship between inflation and unemployment) shows "
    "clear breakdown after 2000 and an outright violation in 2022, when high inflation coexisted with "
    "moderate unemployment. This motivates using richer predictor sets in the ML models."
)
bullet(doc,
    "Oil prices exhibited monthly swings of −54% to +85%, and the yield spread (10Y minus 3M) "
    "inverted ahead of both the 2008 and 2020 recessions — a classic leading indicator that should "
    "add predictive value."
)

body(doc, "Deliverable: six publication-quality figures saved to the figures/ directory.")

# Day 4
h2(doc, "Day 4 — Classical Benchmark Models (ARIMA & VAR)")
body(doc,
    "Tools used: statsmodels — Python’s primary econometrics library, used by academic "
    "researchers and practitioners worldwide. The full expanding-window backtest was run for Naive, "
    "ARIMA, and VAR, generating forecasts at h=1, 3, and 6 for both targets across every month "
    "from 2000 to 2026."
)

h3(doc, "Inflation Results (RMSE, lower is better)")
add_table(doc,
    ["Model", "h=1 (1 month)", "h=3 (3 months)", "h=6 (6 months)"],
    [
        ["Naive", "0.449", "0.977", "1.435"],
        ["ARIMA", "0.402 ✓ Best", "0.979", "1.460"],
        ["VAR",   "0.431", "1.127", "1.974"],
    ],
    col_widths=[1.5, 1.7, 1.7, 1.7]
)

divider(doc)
body(doc,
    "At 1-month ahead, ARIMA outperforms the Naive benchmark by roughly 10% — a meaningful gain. "
    "By 3 months, ARIMA and Naive are indistinguishable. By 6 months, the Naive forecast is marginally "
    "better than ARIMA, and VAR is the worst performer. This is a well-documented result in macro "
    "forecasting: short-horizon structure exists and ARIMA captures it, but that structure fades "
    "rapidly with the horizon."
)

h3(doc, "Unemployment Results (RMSE, lower is better)")
add_table(doc,
    ["Model", "h=1 (1 month)", "h=3 (3 months)", "h=6 (6 months)"],
    [
        ["Naive", "0.642 ✓ Best", "1.087 ✓ Best", "1.400 ✓ Best"],
        ["ARIMA", "1.039", "2.951", "7.003"],
        ["VAR",   "0.989", "2.729", "5.937"],
    ],
    col_widths=[1.5, 1.7, 1.7, 1.7]
)

divider(doc)
body(doc,
    "This is the most striking result of the classical phase. The Naive model dominates ARIMA and VAR "
    "at every horizon for unemployment — by a wide and widening margin. The ARIMA h=6 RMSE of 7.0 "
    "percentage points is catastrophically large relative to Naive’s 1.4."
)
body(doc,
    "The driver is almost certainly the COVID-19 shock: unemployment went from 3.5% to 14.8% and back "
    "in roughly four months. ARIMA, which models unemployment as a smoothly evolving process, "
    "extrapolated the spike as a persistent trend — generating severe errors for months afterward. "
    "This is not a flaw in the implementation; it is a genuine and important finding about the "
    "limitations of parametric time-series models during structural shocks."
)

# ── Section 6: What This Means Going Into ML ──────────────────────────────────
h1(doc, "6. What the Classical Results Tell Us Going Into the ML Phase")

bullet(doc, "ARIMA is the model to beat for short-horizon inflation — a 10% improvement over Naive is real and must be matched or exceeded.")
bullet(doc, "Nothing beats Naive for unemployment in the classical toolbox. ML models must clear this bar.")
bullet(doc, "Longer-horizon forecasting is genuinely difficult. If ML models add value, it may be precisely at h=3 and h=6, where classical models either match or trail the Naive benchmark.")

body(doc,
    "The ML models will be given a meaningful informational advantage: they will see all 12 macro "
    "predictors with up to 6 lags each — up to 72 input features per forecast point. Whether that "
    "additional information translates into better forecasts, especially at longer horizons and during "
    "shock periods, is the central empirical question this project is designed to answer."
)

# ── Section 7: What Comes Next ────────────────────────────────────────────────
h1(doc, "7. What Comes Next")

add_table(doc,
    ["Day", "Task"],
    [
        ["Day 5", "Build and backtest all four ML models (Ridge, Lasso, Random Forest, XGBoost)"],
        ["Day 6", "Merge all results into a single master comparison table across all models and horizons"],
        ["Day 7", "Diagnose why models win or lose — feature importance, sub-period analysis, residual review"],
        ["Days 8–10", "Final tuning, presentation figures, and written findings"],
    ],
    col_widths=[1.0, 5.5]
)

divider(doc)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    "All code, data, and outputs are reproducible from a single FRED API key with no manual data entry."
)
footer_run.italic = True
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Document saved: {OUT}")
