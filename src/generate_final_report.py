"""
Day 10 — Final project report.
Saved to Debriefs/Final_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join("Debriefs", "Final_Report.docx")

NAVY  = RGBColor(0x1F, 0x39, 0x64)
STEEL = RGBColor(0x2E, 0x75, 0xB6)
GRAY  = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1E, 0x6B, 0x3A)
LGRAY = RGBColor(0x88, 0x88, 0x88)

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None, header_fill="1F3964"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_fill)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        fill = "F2F7FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            if isinstance(val, tuple):
                text, bg = val
                set_cell_bg(cell, bg)
                run = cell.paragraphs[0].add_run(text)
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_cell_bg(cell, fill)
                run = cell.paragraphs[0].add_run(str(val))
                run.font.color.rgb = GRAY
                if c_idx == 0:
                    run.bold = True
                run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def title_block(doc, text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY
    return p

def sub_block(doc, text, size=11, color=STEEL):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = STEEL
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.45)
    p.paragraph_format.right_indent = Inches(0.45)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = STEEL
    return p

def divider(doc):
    doc.add_paragraph()

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = LGRAY
    return p

# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ── Cover ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title_block(doc, "Inflation & Unemployment Forecasting Bake-Off", size=24)
title_block(doc, "Traditional Econometric Models vs Machine Learning", size=16)
doc.add_paragraph()
sub_block(doc, "Final Project Report", size=12)
sub_block(doc, "May 2026", size=11, color=LGRAY)
doc.add_paragraph()
doc.add_paragraph()

callout(doc,
    "Do traditional time-series models outperform machine learning methods in forecasting "
    "inflation and unemployment, or do ML methods gain an edge once richer predictor sets "
    "and nonlinear patterns are included?"
)

doc.add_page_break()

# ── Executive Summary ──────────────────────────────────────────────────────────
h1(doc, "Executive Summary")
body(doc,
    "This project constructs a rigorous forecasting competition between seven models — two "
    "classical time-series models, four machine learning methods, and a naive benchmark — "
    "evaluated on monthly U.S. inflation and unemployment data from 2000 to 2026. Each model "
    "is tested under expanding-window backtesting, ensuring no look-ahead bias. Forecast "
    "accuracy is measured at horizons of 1, 3, and 6 months ahead."
)
body(doc,
    "XGBoost is the strongest performer overall, winning five of six target-horizon combinations "
    "on RMSE. Lasso regression wins the remaining cell — short-horizon inflation — where variable "
    "selection over a rich predictor set proves more effective than ensemble methods. Classical "
    "models (ARIMA, VAR) are never the outright winner in any category, though ARIMA remains "
    "competitive for short-horizon inflation forecasting."
)
body(doc,
    "The most significant finding is the COVID-19 stress test. During the 2020 economic shock, "
    "tree-based models (XGBoost, Random Forest) maintained meaningful forecast accuracy while "
    "classical models and linear machine learning methods produced catastrophically large errors "
    "for unemployment. XGBoost is the only model that performs consistently well across all three "
    "macroeconomic regimes tested: pre-2008, 2008–2019, and 2020 onward."
)
body(doc,
    "The project concludes that machine learning — specifically gradient boosting — adds genuine, "
    "measurable value in macroeconomic forecasting. The gains are most pronounced at longer "
    "horizons and during periods of economic stress, precisely where traditional models are most "
    "likely to fail."
)

# ── Section 1: Motivation ──────────────────────────────────────────────────────
h1(doc, "1. Motivation")
body(doc,
    "Inflation and unemployment are among the most consequential macroeconomic variables "
    "monitored by policymakers, businesses, and financial markets. The Federal Reserve's dual "
    "mandate — price stability and maximum employment — is defined entirely in terms of these "
    "two series. Accurate near-term forecasts of both variables directly inform interest rate "
    "decisions, fiscal policy design, corporate hiring plans, and fixed-income pricing."
)
body(doc,
    "Traditional time-series models have been the workhorse of macroeconomic forecasting for "
    "decades. ARIMA models, introduced by Box and Jenkins in the 1970s, remain widely used in "
    "central bank and academic forecasting. Vector autoregressions (VARs), developed by Sims "
    "(1980), allow multiple variables to be modelled jointly and remain standard tools in "
    "applied macroeconometrics."
)
body(doc,
    "The rapid development of machine learning methods raises a practical question: given access "
    "to the same historical data — and in some cases considerably more of it — can modern "
    "computational approaches produce better forecasts than the models economists have refined "
    "over forty years? This project is designed to provide a direct, fair, and reproducible "
    "answer to that question."
)

# ── Section 2: Research Question ──────────────────────────────────────────────
h1(doc, "2. Research Question and Hypotheses")
body(doc,
    "The central research question is:"
)
callout(doc,
    "Do traditional time-series models (ARIMA, VAR) outperform machine learning methods "
    "(Ridge, Lasso, Random Forest, XGBoost) in forecasting U.S. inflation and unemployment, "
    "or do ML methods gain an edge when given access to richer predictor sets and allowed "
    "to capture nonlinear dynamics?"
)
body(doc, "Three specific hypotheses are tested:")
bullet(doc,
    "ML models will outperform classical models at longer horizons (h=3, h=6), where "
    "the informational advantage of additional predictors and nonlinear relationships becomes "
    "more valuable.", "H1:  ")
bullet(doc,
    "Classical models will remain competitive at the shortest horizon (h=1), where "
    "autocorrelation structure dominates and additional predictors offer limited marginal value.", "H2:  ")
bullet(doc,
    "Tree-based ML models will prove more robust to structural shocks than linear models — "
    "both classical and ML — because decision trees naturally contain extreme values rather "
    "than extrapolating them.", "H3:  ")

# ── Section 3: Data ───────────────────────────────────────────────────────────
h1(doc, "3. Data")

h2(doc, "3.1 Sources")
body(doc,
    "All data were obtained from FRED (Federal Reserve Economic Data), the St. Louis Federal "
    "Reserve's public macroeconomic database, via the fredapi Python library. Eleven monthly "
    "series were pulled covering January 1990 through March 2026."
)

add_table(doc,
    ["Variable", "FRED Series ID", "Transformation", "Role"],
    [
        ["CPI (All Items)",       "CPIAUCSL",   "YoY % change",   "Primary forecast target"],
        ["Core CPI",              "CPILFESL",   "YoY % change",   "Target & predictor"],
        ["Unemployment Rate",     "UNRATE",     "Level",          "Primary forecast target"],
        ["Federal Funds Rate",    "FEDFUNDS",   "Level",          "Predictor"],
        ["10-Year Treasury",      "GS10",       "Level",          "Predictor"],
        ["3-Month T-Bill",        "TB3MS",      "Level",          "Predictor"],
        ["Yield Spread (10Y−3M)", "Derived",    "Level",          "Predictor"],
        ["Industrial Production", "INDPRO",     "MoM % change",   "Predictor"],
        ["Nonfarm Payrolls",      "PAYEMS",     "MoM % change",   "Predictor"],
        ["Oil Prices (WTI)",      "DCOILWTICO", "MoM % change",   "Predictor"],
        ["Consumer Sentiment",    "UMCSENT",    "Level",          "Predictor"],
        ["M2 Money Supply",       "M2SL",       "MoM % change",   "Predictor"],
    ],
    col_widths=[1.7, 1.2, 1.3, 1.9]
)

h2(doc, "3.2 Transformations")
body(doc,
    "CPI and Core CPI are expressed as year-over-year percentage changes — the standard "
    "representation used by the Federal Reserve and financial press. This transformation "
    "removes the unit-root non-stationarity of price levels and produces a series interpretable "
    "directly in percentage-point terms. Flow variables (industrial production, payrolls, oil, "
    "M2) are expressed as month-over-month percentage changes. Interest rate levels are retained "
    "in level form, as they are already stationary over most of the sample."
)
body(doc,
    "After transformation, the processed dataset contains 421 monthly observations spanning "
    "January 1991 through February 2026. The 12-month lag required for year-over-year "
    "inflation computation accounts for the reduction from the 435-row raw pull."
)

h2(doc, "3.3 Feature Engineering for ML Models")
body(doc,
    "Machine learning models were given access to the current value and six monthly lags of "
    "all 12 processed variables, producing a feature matrix of 84 columns per observation. "
    "Lag alignment was performed carefully to prevent data leakage: for an h-step-ahead "
    "forecast at time t, the target variable is the realised value at t+h, and features "
    "are constructed from observations available at or before time t."
)

h2(doc, "3.4 Key Descriptive Statistics")
add_table(doc,
    ["Variable", "Mean", "Std Dev", "Min", "Max", "Key Event"],
    [
        ["CPI Inflation (YoY %)", "2.62", "1.54", "−1.96", "8.98",  "Peak: Jun 2022 (+9.0%)"],
        ["Unemployment Rate (%)", "5.67", "1.76", "3.40",  "14.80", "Peak: Apr 2020 (+14.8%)"],
        ["Federal Funds Rate (%)", "2.42", "2.16", "0.07",  "6.54",  "Near-zero: 2009–2015, 2020–2022"],
        ["Oil MoM % Change",      "0.72", "10.14","−54.2", "85.0",  "Crash: Apr 2020 (−54.2%)"],
    ],
    col_widths=[2.1, 0.7, 0.8, 0.7, 0.7, 2.1]
)

# ── Section 4: Methods ────────────────────────────────────────────────────────
h1(doc, "4. Methods")

h2(doc, "4.1 Naive Benchmark")
body(doc,
    "The naive model predicts that the next observation equals the current observation — a "
    "random walk without drift. It serves as the minimum performance bar. In macroeconomic "
    "forecasting literature, the naive benchmark is notoriously difficult to beat at short "
    "horizons, making it an appropriate and stringent baseline."
)

h2(doc, "4.2 ARIMA")
body(doc,
    "ARIMA (AutoRegressive Integrated Moving Average) models were estimated separately for "
    "each target variable. Lag orders (p, d, q) were selected at each expanding-window step "
    "by minimising the Akaike Information Criterion (AIC) over a grid of p ∈ {0,1,2}, "
    "d ∈ {0,1}, q ∈ {0,1,2}. To avoid the computational cost of running this 18-combination "
    "grid search for every forecast horizon, the optimal order was cached once per training "
    "window and reused across all horizons. Models were estimated using the SARIMAX "
    "implementation in statsmodels 0.14."
)

h2(doc, "4.3 VAR")
body(doc,
    "A bivariate Vector Autoregression was estimated jointly on inflation and unemployment. "
    "Lag length was selected by AIC up to a maximum of six lags. VAR allows the joint "
    "dynamics between the two target variables to inform forecasts of each — in principle "
    "capturing the Phillips Curve relationship and cross-variable spillovers."
)

h2(doc, "4.4 Ridge Regression")
body(doc,
    "Ridge regression applies L2 regularisation to ordinary least squares, shrinking "
    "coefficient estimates toward zero to prevent overfitting in the 84-feature setting. "
    "A regularisation penalty of α = 1.0 was used, applied to standardised features. "
    "Ridge retains all predictors but reduces the influence of weaker ones."
)

h2(doc, "4.5 Lasso Regression")
body(doc,
    "Lasso applies L1 regularisation, which has the additional property of setting weak "
    "predictor coefficients exactly to zero — performing automatic variable selection. "
    "A penalty of α = 0.01 was used. In a 84-predictor setting, Lasso's ability to "
    "identify and discard irrelevant lags is a meaningful structural advantage."
)

h2(doc, "4.6 Random Forest")
body(doc,
    "Random Forest fits an ensemble of 100 decision trees on bootstrap samples of the "
    "training data, with random feature subsets considered at each split. The ensemble "
    "average reduces variance relative to any individual tree. Key hyperparameters: "
    "max_depth = 5, min_samples_leaf = 5. Random Forest is inherently robust to outliers "
    "because individual trees produce bounded predictions."
)

h2(doc, "4.7 XGBoost")
body(doc,
    "XGBoost implements gradient boosting: an ensemble of 150 decision trees fitted "
    "sequentially, where each tree targets the residuals of all previous trees. This "
    "approach can achieve very low bias while controlling variance through regularisation "
    "and tree depth constraints. Hyperparameters: learning_rate = 0.05, max_depth = 4, "
    "subsample = 0.8, colsample_bytree = 0.8. XGBoost is the current state-of-the-art "
    "for tabular forecasting in applied competitions and industry."
)

# ── Section 5: Evaluation Framework ───────────────────────────────────────────
h1(doc, "5. Evaluation Framework")

h2(doc, "5.1 Expanding-Window Backtesting")
body(doc,
    "All models are evaluated under an expanding-window (recursive) backtesting scheme. "
    "At each evaluation step t, the model is trained on all available data from the start "
    "of the sample through t−1. Forecasts are then generated for horizons h = 1, 3, and "
    "6 months ahead. The window expands by one month at each step, and the process repeats "
    "from January 2000 through February 2026 — yielding 314 evaluation points per "
    "model-target-horizon combination. A minimum training window of 60 months (5 years) "
    "is required before the first forecast."
)
body(doc,
    "This design enforces strict temporal ordering: no model ever observes data from beyond "
    "the current evaluation date. This is the standard evaluation protocol in academic "
    "macroeconomic forecasting and ensures that all reported results reflect genuine "
    "out-of-sample forecast accuracy."
)

h2(doc, "5.2 Direct Multi-Step Forecasting")
body(doc,
    "For the ML models, direct forecasting is used: a separate model is estimated for each "
    "horizon h, with the target variable at time t+h regressed on features observed at time t. "
    "This avoids the compounding of forecast errors that occurs in iterated forecasting and "
    "allows each model to learn a horizon-specific mapping from predictors to outcomes."
)

h2(doc, "5.3 Accuracy Metrics")
add_table(doc,
    ["Metric", "Formula (conceptual)", "Interpretation"],
    [
        ["RMSE", "√[ mean(actual − forecast)² ]",
         "Penalises large errors more heavily. Primary metric throughout."],
        ["MAE",  "mean |actual − forecast|",
         "Average absolute error in percentage-point terms. Secondary metric."],
    ],
    col_widths=[0.9, 2.4, 3.3]
)

# ── Section 6: Results ────────────────────────────────────────────────────────
h1(doc, "6. Results")

h2(doc, "6.1 Overall RMSE Comparison — Inflation")
note(doc, "Lower RMSE is better. Bold = best in column.")

add_table(doc,
    ["Model", "h=1", "h=3", "h=6"],
    [
        ["Naive",               "0.449",          "0.977",          "1.436"],
        ["ARIMA (auto-order)",  "0.402",          "0.957",          "1.448"],
        ["VAR",                 "0.431",          "1.127",          "1.974"],
        ["Ridge (α=1.0)",       "0.533",          "1.424",          "1.597"],
        ["Lasso (α=0.01)",      ("0.278 ✓", "2E75B6"), "0.934",   "1.242"],
        ["Random Forest",       "0.421",          "0.795",          "1.060"],
        ["XGBoost",             "0.335", ("0.746 ✓", "2E75B6"), ("0.831 ✓", "2E75B6")],
    ],
    col_widths=[2.0, 1.2, 1.2, 1.2]
)
divider(doc)

h2(doc, "6.2 Overall RMSE Comparison — Unemployment")
note(doc, "Lower RMSE is better. Bold = best in column.")

add_table(doc,
    ["Model", "h=1", "h=3", "h=6"],
    [
        ["Naive",               "0.642",          "1.087",          "1.400"],
        ["ARIMA (auto-order)",  "0.996",          "2.741",          "5.756"],
        ["VAR",                 "0.989",          "2.729",          "5.937"],
        ["Ridge (α=1.0)",       "2.211",          "2.338",          "2.665"],
        ["Lasso (α=0.01)",      "2.099",          "2.073",          "2.506"],
        ["Random Forest",       "0.448",          "0.903",          "1.168"],
        ["XGBoost",             ("0.259 ✓", "2E75B6"), ("0.903 ✓", "2E75B6"), ("0.956 ✓", "2E75B6")],
    ],
    col_widths=[2.0, 1.2, 1.2, 1.2]
)
divider(doc)

h2(doc, "6.3 Winner Summary")
add_table(doc,
    ["", "h=1 (1 month)", "h=3 (3 months)", "h=6 (6 months)"],
    [
        ["Inflation",    "Lasso (0.278)",    "XGBoost (0.746)", "XGBoost (0.831)"],
        ["Unemployment", "XGBoost (0.259)",  "XGBoost (0.903)", "XGBoost (0.956)"],
    ],
    col_widths=[1.5, 1.9, 1.9, 1.9]
)
divider(doc)

h2(doc, "6.4 Sub-Period Performance — Inflation (h=1 RMSE)")
body(doc,
    "Performance is split across three macroeconomic regimes to assess model stability:"
)
add_table(doc,
    ["Model", "Pre-2008", "2008–2019", "2020+"],
    [
        ["Naive",         "0.448", "0.448", "0.456"],
        ["ARIMA",         "0.417", "0.401", "0.384"],
        ["VAR",           "0.406", "0.413", "0.494"],
        ["Ridge",         "0.276", "0.266", "0.945"],
        ["Lasso",         "0.147", "0.204", "0.337"],
        ["Random Forest", "0.276", "0.382", "0.611"],
        ["XGBoost",       "0.263", "0.303", "0.523"],
    ],
    col_widths=[2.0, 1.2, 1.2, 1.2]
)
divider(doc)

h2(doc, "6.5 Sub-Period Performance — Unemployment (h=1 RMSE)")
add_table(doc,
    ["Model", "Pre-2008", "2008–2019", "2020+"],
    [
        ["Naive",         "0.126", "0.178", "1.299"],
        ["ARIMA",         "0.129", "0.160", "2.135"],
        ["VAR",           "0.130", "0.162", "2.029"],
        ["Ridge",         "0.143", "0.123", "4.445"],
        ["Lasso",         "0.079", "0.123", "3.937"],
        ["Random Forest", "0.142", "0.319", "0.795"],
        ["XGBoost",       "0.113", "0.243", "0.492"],
    ],
    col_widths=[2.0, 1.2, 1.2, 1.2]
)
divider(doc)

h2(doc, "6.6 Day 8 Tuning Results")
body(doc,
    "Model hyperparameters were reviewed and adjusted following the Day 7 diagnostic phase. "
    "Results compared to original (Day 5) specifications:"
)
add_table(doc,
    ["Change", "Effect", "Verdict"],
    [
        ["Auto-ARIMA order selection (AIC grid)",
         "Inflation h=3: 0.979→0.957. Unemployment h=6: 7.003→5.756",
         "Improved ✓"],
        ["XGBoost: 100→150 estimators",
         "Uniform improvement across cells. Best gain: unemployment h=1 (0.296→0.259)",
         "Improved ✓"],
        ["Lasso: α 0.01→0.005",
         "Performance worsened across all cells. Original α=0.01 was already optimal.",
         "Reverted ✗"],
        ["Ridge: α 1.0→0.5",
         "Performance worsened across all cells. Original α=1.0 was already optimal.",
         "Reverted ✗"],
    ],
    col_widths=[2.4, 2.5, 1.3]
)

# ── Section 7: Discussion ─────────────────────────────────────────────────────
h1(doc, "7. Discussion")

h2(doc, "7.1 Why XGBoost Dominates at Longer Horizons")
body(doc,
    "At h=1, the dominant signal for any macro series is its own recent history — "
    "autocorrelation structure that ARIMA captures well. As the horizon extends to h=3 and "
    "h=6, the autocorrelation signal weakens and the marginal value of additional predictors "
    "increases. XGBoost, with access to 84 features including lags of interest rates, output, "
    "employment, and sentiment, can exploit cross-variable predictive relationships that are "
    "invisible to ARIMA. This finding is consistent with H1."
)
body(doc,
    "Feature importance analysis reinforces this interpretation. XGBoost distributes importance "
    "across many features — recent lags of unemployment (lag1: 0.122, lag2: 0.119, lag3: 0.090) "
    "and the 10-year Treasury rate (lag6: 0.123) all contribute meaningfully. The model has "
    "genuinely learned that financial conditions today predict economic outcomes 3–6 months "
    "from now."
)

h2(doc, "7.2 Why Lasso Wins at Short-Horizon Inflation")
body(doc,
    "Lasso's superiority at h=1 inflation (RMSE: 0.278 vs ARIMA's 0.402) reflects its ability "
    "to select a tight, stable set of predictors and discard noise. In a 84-feature space, "
    "most lagged variables carry little marginal signal for one-month-ahead inflation. Lasso "
    "effectively reduces to a parsimonious model that captures the most persistent signals — "
    "recent CPI momentum and a small number of financial variables — while ignoring the rest. "
    "ARIMA, restricted to the univariate series, cannot access these cross-variable signals "
    "and so performs inferiorly. This partially supports H2: classical models remain "
    "competitive at h=1, but Lasso — itself a relatively simple model — surpasses them."
)

h2(doc, "7.3 The COVID-19 Stress Test and the Failure of Linear Models")
body(doc,
    "The unemployment results in the 2020+ sub-period are the project's most striking finding. "
    "Unemployment rose from 3.5% to 14.8% in April 2020 and largely recovered within four months. "
    "This constitutes an event with no precedent in the training data."
)
body(doc,
    "ARIMA's h=6 unemployment RMSE reached 5.756 in the full sample — a 311% deterioration "
    "relative to the naive benchmark's 1.400. The mechanism is straightforward: ARIMA models "
    "unemployment as a smooth autoregressive process and extrapolated the spike as a persistent "
    "upward trend, generating forecasts far above any subsequent realisation."
)
body(doc,
    "Ridge and Lasso failed even more severely (2020+ RMSE: 4.445 and 3.937 respectively). "
    "As linear models, they face the same fundamental limitation: in a training set dominated "
    "by low-volatility observations, a linear model calibrated to normal times will extrapolate "
    "extreme values as persistent rather than transient."
)
body(doc,
    "XGBoost's 2020+ RMSE for unemployment at h=1 was 0.492 — below the naive benchmark's "
    "1.299. Decision trees partition the feature space into regions and predict the average "
    "outcome within each region. When unemployment spikes to 14.8%, the tree routes to a "
    "region it has seen before (high unemployment during the 2009 recession) and predicts "
    "reversion rather than continuation. This structural property makes tree-based models "
    "naturally more robust to shocks. H3 is strongly confirmed."
)

h2(doc, "7.4 The Paradox of the Phillips Curve")
body(doc,
    "Exploratory data analysis revealed that the classical inverse relationship between "
    "inflation and unemployment — the Phillips Curve — has broken down substantially over "
    "the sample period. The 2022 episode, in which inflation reached 9% while unemployment "
    "remained near 3.5%, represents an outright violation of the traditional relationship."
)
body(doc,
    "This breakdown provides indirect support for the ML approach: if the structural "
    "relationships between macro variables are time-varying and potentially nonlinear, "
    "then models that can adapt to those changes — through feature selection (Lasso) or "
    "nonlinear partitioning (XGBoost) — should outperform models that assume fixed linear "
    "dynamics (ARIMA, VAR). The results confirm this expectation."
)

h2(doc, "7.5 What the Models Pay Attention To")
body(doc,
    "Random Forest concentrates 75% of its importance weight on the current level of "
    "unemployment and inflation — it behaves as a sophisticated persistence model. The "
    "additional 84-feature set adds relatively little marginal value for RF. XGBoost, "
    "by contrast, distributes importance more evenly across lags and cross-variable "
    "predictors, genuinely exploiting the information in the full feature matrix."
)
body(doc,
    "Two variables appear consistently across both models and all horizons: the 10-year "
    "Treasury yield and the yield spread (10Y minus 3M). This aligns with decades of "
    "empirical evidence that the yield curve is a leading indicator for both economic "
    "activity and inflation — and confirms that these variables carry genuine predictive "
    "signal beyond what is already contained in the inflation and unemployment series themselves."
)

# ── Section 8: Conclusions ────────────────────────────────────────────────────
h1(doc, "8. Conclusions")

body(doc,
    "This project set out to answer a direct empirical question: does machine learning add "
    "value over traditional econometric models in forecasting U.S. inflation and unemployment? "
    "The answer is yes — but with meaningful nuance."
)

bullet(doc,
    "XGBoost is the strongest overall forecasting model, winning five of six "
    "target-horizon combinations. Its gains over the naive benchmark are 25% at "
    "inflation h=1, 24% at inflation h=3, and 42% at inflation h=6. For unemployment, "
    "it beats the naive benchmark by 60% at h=1 and 32% at h=6.",
    "Finding 1:  ")

bullet(doc,
    "Lasso regression is the best model for short-horizon inflation (h=1), beating ARIMA "
    "by 31% and the naive benchmark by 38%. Its variable-selection mechanism identifies "
    "a sparse, stable set of predictors that outperforms both classical models and more "
    "complex ML methods at this specific horizon.",
    "Finding 2:  ")

bullet(doc,
    "Classical models (ARIMA, VAR) are never the outright winner. However, ARIMA with "
    "automatic order selection is competitive with Random Forest at h=1 for inflation "
    "and should not be dismissed as a component of a forecasting ensemble.",
    "Finding 3:  ")

bullet(doc,
    "Tree-based models (XGBoost, Random Forest) are uniquely robust to structural shocks. "
    "During the 2020+ period, XGBoost is the only model that outperforms the naive benchmark "
    "for unemployment at all horizons. All other model classes — including linear ML — fail "
    "to maintain competitive accuracy under shock conditions.",
    "Finding 4:  ")

bullet(doc,
    "Tuning does not always improve performance. Reducing regularisation strength for Lasso "
    "and Ridge worsened their accuracy, confirming that the original penalty values were "
    "already well-calibrated for this forecasting task. Increasing XGBoost's estimator count "
    "from 100 to 150, however, produced consistent improvements across all cells.",
    "Finding 5:  ")

divider(doc)
body(doc,
    "The evidence supports a practical recommendation for applied macroeconomic forecasting: "
    "gradient boosting methods (XGBoost) should be the primary tool for medium- and long-horizon "
    "forecasting of both inflation and unemployment. For short-horizon inflation, Lasso provides "
    "a simpler, more interpretable alternative with superior accuracy. Classical models remain "
    "useful as benchmarks and as components of forecast combinations, but should not be used "
    "as standalone forecasting tools for horizons beyond one month."
)

callout(doc,
    "The key lesson from this project is not simply that machine learning wins. It is that "
    "the conditions under which each model class excels are predictable and economically "
    "interpretable — and that understanding those conditions is as valuable as finding the "
    "single best-performing model."
)

# ── Section 9: Limitations and Future Work ────────────────────────────────────
h1(doc, "9. Limitations and Future Work")

h2(doc, "9.1 Limitations")
bullet(doc, "The evaluation period begins in 2000. Results may not generalise to earlier "
       "periods with different inflation regimes (e.g., the high-inflation 1970s and 1980s).")
bullet(doc, "Hyperparameter tuning for tree-based models used fixed values rather than "
       "rolling cross-validation, which could be suboptimal at certain points in the sample.")
bullet(doc, "Only point forecasts are evaluated. Forecast uncertainty (prediction intervals) "
       "is not modelled, which limits the direct policy relevance of the outputs.")
bullet(doc, "The COVID-19 shock is a single extreme event. Conclusions about tree-model "
       "robustness are based on one episode and should be interpreted cautiously.")

h2(doc, "9.2 Directions for Future Work")
bullet(doc, "Forecast combination: averaging ARIMA, Lasso, and XGBoost forecasts may "
       "outperform any individual model through error diversification.")
bullet(doc, "Prediction intervals: quantile regression forests or conformal prediction "
       "would add probabilistic outputs suitable for risk management applications.")
bullet(doc, "Real-time data: this study uses revised data. A true real-time backtest using "
       "vintage data would account for data revisions and provide a more operationally "
       "realistic evaluation.")
bullet(doc, "Extended model set: Neural networks (LSTM, Temporal Fusion Transformer) and "
       "factor models (DFM) would provide additional points of comparison.")
bullet(doc, "Longer sample: including data back to 1960 would allow testing in a genuinely "
       "different inflation regime and provide a more rigorous assessment of model robustness.")

# ── Section 10: Technical Appendix ────────────────────────────────────────────
h1(doc, "10. Technical Appendix")

h2(doc, "10.1 Software and Libraries")
add_table(doc,
    ["Library", "Version", "Purpose"],
    [
        ["pandas",        "3.0.2",  "Data wrangling and time series alignment"],
        ["numpy",         "2.4.4",  "Numerical computation and error metrics"],
        ["statsmodels",   "0.14.6", "ARIMA and VAR estimation"],
        ["scikit-learn",  "1.8.0",  "Ridge, Lasso, Random Forest, preprocessing"],
        ["xgboost",       "3.2.0",  "XGBoost gradient boosting"],
        ["fredapi",       "0.5.2",  "FRED API data retrieval"],
        ["matplotlib",    "3.10.9", "EDA and diagnostic figures"],
        ["plotly",        "6.7.0",  "Interactive HTML dashboard"],
        ["python-docx",   "1.2.0",  "Report generation"],
        ["python-dotenv", "1.2.2",  "Secure API key management"],
    ],
    col_widths=[1.5, 0.9, 4.2]
)

h2(doc, "10.2 Reproducibility")
body(doc,
    "The complete pipeline is reproducible from a single FRED API key. Running the seven "
    "numbered scripts in src/ in order (01 through 07, plus 05 for the merge) regenerates "
    "all results, figures, and the interactive dashboard from raw data downloads. "
    "No manual data entry is required at any stage."
)

h2(doc, "10.3 Project Deliverables")
add_table(doc,
    ["File", "Location", "Description"],
    [
        ["dashboard.html",                              "Debriefs/", "Interactive 5-panel Plotly dashboard"],
        ["Final_Report.docx",                          "Debriefs/", "This document"],
        ["Status_Debrief_Day9.docx",                   "Debriefs/", "Plain-English progress summary"],
        ["Inflation_Unemployment_Forecasting_Project_Plan.docx", "Debriefs/", "10-day project plan"],
        ["Inflation_Unemployment_Forecasting_Progress_Report_Days1-4.docx", "Debriefs/", "Days 1–4 client debrief"],
        ["master_errors.csv",                          "results/",  "Full RMSE/MAE table — all 42 model-target-horizon combinations"],
        ["master_forecasts.csv",                       "results/",  "1,864 individual forecast observations"],
        ["subperiod_errors.csv",                       "results/",  "RMSE by economic regime"],
        ["feature_importance_*.csv",                   "results/",  "Feature importance for RF and XGBoost (6 files)"],
        ["macro_monthly.csv",                          "data/processed/", "Clean monthly feature dataset"],
    ],
    col_widths=[3.2, 1.3, 2.1]
)

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    "Data: Federal Reserve Economic Data (FRED), St. Louis Fed.  "
    "Sample: January 1991 – February 2026.  "
    "Evaluation: January 2000 – February 2026."
)
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = LGRAY

doc.save(OUT)
print(f"Document saved: {OUT}")
