"""
Generates a plain-English project explainer Word document.
Saved to Debriefs/Project_Explainer.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join("Debriefs", "Project_Explainer.docx")

NAVY  = RGBColor(0x1F, 0x39, 0x64)
STEEL = RGBColor(0x2E, 0x75, 0xB6)
GRAY  = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x37, 0x86, 0x44)
AMBER = RGBColor(0xBF, 0x87, 0x00)
LGRAY = RGBColor(0xD0, 0xD0, 0xD0)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = color or NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = color or STEEL
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = color or GRAY
    return p

def body(doc, text, color=None, italic=False, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or GRAY
    run.italic = italic
    run.bold = bold
    return p

def bullet(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or GRAY
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg="1F3964", bold_winner=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        highlight = bold_winner and row[0] == bold_winner
        bg = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY
            if highlight:
                run.bold = True
                run.font.color.rgb = GREEN

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

# ─── BUILD DOCUMENT ──────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ── Cover ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Inflation & Unemployment Forecasting Bake-Off")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Plain-English Guide: What We Did and What We Found")
    r.font.size = Pt(13)
    r.font.color.rgb = STEEL
    r.italic = True

    doc.add_paragraph()
    divider(doc)
    doc.add_paragraph()

    # ── Part 1: Hypothesis ────────────────────────────────────────────────────
    heading(doc, "Part 1 — The Hypothesis")
    body(doc,
        "The central question of this project was: Are traditional economic models better at predicting "
        "inflation and unemployment, or do modern machine learning models do a better job?"
    )
    body(doc,
        "Economists have used classical statistical models for decades. But in recent years, machine "
        "learning — the same family of techniques behind recommendation algorithms and fraud detection — "
        "has started showing up in economic research. The intuition behind testing machine learning here "
        "is that the economy is complicated and nonlinear. Inflation doesn't just go up when one thing "
        "goes up. It responds to dozens of signals at once, and those relationships change over time. "
        "Classical models assume relatively simple, fixed relationships. Machine learning models can "
        "discover more complex, shifting patterns automatically."
    )
    body(doc,
        "So the hypothesis was: machine learning models will outperform classical models, especially at "
        "longer forecast horizons and during unusual economic episodes like the COVID-19 pandemic.",
        italic=True
    )

    divider(doc)

    # ── Part 2: Variables ─────────────────────────────────────────────────────
    heading(doc, "Part 2 — The Variables")
    body(doc,
        "We used 12 economic indicators pulled from FRED (the Federal Reserve's public data library), "
        "covering data from January 1990 through early 2026. Here is what each one measures and why "
        "it matters for forecasting."
    )

    heading(doc, "What We Were Trying to Predict (The Targets)", level=2)
    add_table(doc,
        headers=["Variable", "What It Measures"],
        rows=[
            ["CPI Inflation (Year-over-Year %)",
             "How much more expensive goods and services are compared to the same month one year ago. "
             "This is the headline inflation number you hear in the news."],
            ["Unemployment Rate (%)",
             "The share of people who want a job but don't have one. A simple but powerful gauge "
             "of economic health."],
        ],
        col_widths=[2.2, 4.1],
        header_bg="1F3964",
    )

    heading(doc, "The Predictors (Signals the Models Used to Make Forecasts)", level=2)
    predictor_rows = [
        ["Core CPI Inflation (YoY %)",
         "Inflation excluding food and energy prices",
         "A stable signal of underlying price pressure; less noisy than headline CPI"],
        ["Federal Funds Rate (%)",
         "The interest rate the Fed sets for overnight lending between banks",
         "When the Fed raises rates, it slows the economy, cooling inflation and raising unemployment"],
        ["10-Year Treasury Yield (%)",
         "The interest rate on 10-year government bonds",
         "Reflects long-term growth and inflation expectations"],
        ["3-Month Treasury Bill Rate (%)",
         "Short-term interest rate",
         "Used to build the yield spread; also signals credit conditions"],
        ["Yield Spread (10Y minus 3M, %)",
         "The gap between long and short-term interest rates",
         "A famous recession predictor — when short rates exceed long rates, recession often follows"],
        ["Industrial Production (monthly % change)",
         "How much factories, mines, and utilities produced",
         "A real-time gauge of economic activity; leads unemployment"],
        ["Nonfarm Payrolls (monthly % change)",
         "How many jobs were added or lost outside of farming",
         "One of the most watched economic releases; directly tied to unemployment"],
        ["Oil Prices — WTI (monthly % change)",
         "Price of crude oil",
         "Oil prices feed directly into inflation; shocks (like 2022) are often inflation-driving"],
        ["Consumer Sentiment",
         "A survey of how optimistic consumers feel about the economy",
         "Forward-looking; when sentiment drops, spending slows and recession risk rises"],
        ["M2 Money Supply (monthly % change)",
         "Total money circulating in the economy (cash, savings, etc.)",
         "Too much money chasing goods causes inflation — the core of monetarist theory"],
    ]
    add_table(doc,
        headers=["Variable", "What It Measures", "Why It Helps Forecasting"],
        rows=predictor_rows,
        col_widths=[1.8, 2.0, 2.5],
        header_bg="2E75B6",
    )

    body(doc,
        "For every one of these 12 variables, we also created 6 time-lagged copies — meaning the models "
        "could see not just today's value but also values from 1, 2, 3, 4, 5, and 6 months prior. That "
        "produced an 84-column feature set that the machine learning models used to learn patterns from history."
    )

    divider(doc)

    # ── Part 3: How We Built It ───────────────────────────────────────────────
    heading(doc, "Part 3 — How We Built the Program")
    body(doc,
        "The project ran over 10 days, with a different script built each day. Here is how the pipeline "
        "flowed, step by step."
    )

    steps = [
        ("Day 1 — Data Collection",
         "We connected to FRED's API (a data feed the Federal Reserve provides for free) and downloaded "
         "all 12 series going back to January 1990. The script automatically cleaned, aligned, and "
         "transformed the raw data — computing year-over-year percent changes for CPI, percentage growth "
         "rates for production and payrolls, and keeping interest rates as plain levels. The final output "
         "was a single clean spreadsheet with 421 monthly rows and 12 columns."),
        ("Day 2 — Exploratory Analysis",
         "Before building any models, we visualized the data to understand its structure. We looked at "
         "how each variable moved over time, how they were correlated with each other, and whether "
         "inflation and unemployment moved together (they often move in opposite directions — low "
         "unemployment tends to push inflation up). This step prevents going into modeling blind."),
        ("Day 3 — Classical Models: ARIMA and VAR",
         "ARIMA (AutoRegressive Integrated Moving Average) is a workhorse of economic forecasting. It "
         "looks at a single variable's own past values and past errors to predict its next value — think "
         "of it as: 'Inflation last month was X, the month before was Y, so my best guess for next month "
         "is...' We tested 18 different parameter combinations for each model and selected the best. "
         "VAR (Vector AutoRegression) is ARIMA's more sophisticated cousin: instead of one variable in "
         "isolation, it looks at all 12 simultaneously and captures how they influence each other — for "
         "example, that rising Fed rates tend to reduce inflation with some lag. Both models used an "
         "expanding-window backtest: train on all data up to a given month, forecast 1, 3, and 6 months "
         "ahead, move one month forward, repeat. This ran 313 rounds from January 2000 onward."),
        ("Day 4 — Machine Learning Models",
         "Four ML models were trained using the same expanding-window approach. Ridge Regression: a "
         "linear model that uses all 84 features but applies a mathematical penalty to prevent any single "
         "feature from dominating. Lasso Regression: similar to Ridge, but more aggressive — it can "
         "completely ignore features it finds unhelpful, functioning as an automatic variable selector. "
         "Random Forest: an ensemble of 100 decision trees, each trained on a random subset of data, "
         "voting on the final prediction — handles nonlinear relationships well. XGBoost: a more "
         "sophisticated tree method where trees are built sequentially, each correcting the errors of the "
         "previous one — often the strongest performer on structured data."),
        ("Day 5 — Merging Results",
         "All seven models' forecasts were merged into a single master spreadsheet. We also included a "
         "Naive benchmark — simply predicting that next month's value will equal this month's value. "
         "This is the minimum bar every model has to clear to be considered useful."),
        ("Day 6 — Diagnostics",
         "We dug deeper: residual analysis checked whether model errors were random (good) or patterned "
         "(a sign something is being missed). Feature importance extracted which of the 84 inputs mattered "
         "most — recent lags of inflation itself and unemployment came out on top, followed by the yield "
         "spread and oil prices. Sub-period analysis split results into three eras: pre-2008 (Great "
         "Moderation), 2008–2019 (post-financial-crisis), and 2020+ (COVID era)."),
        ("Days 7–9 — Visualization and Dashboard",
         "An interactive HTML dashboard was built using Plotly. Six panels cover forecast accuracy, "
         "actual vs. predicted time series, model comparison bar charts, sub-period breakdown, forecast "
         "errors over time, and feature importance. Dropdown menus allow switching between targets, "
         "models, and horizons without running any code."),
        ("Day 10 — Final Report",
         "A formal Word document summarizing the full project: methodology, all results tables, tuning "
         "experiments, sub-period findings, and limitations."),
    ]

    for title, desc in steps:
        heading(doc, title, level=2)
        body(doc, desc)

    divider(doc)

    # ── Part 4: Results ───────────────────────────────────────────────────────
    heading(doc, "Part 4 — The Results")
    body(doc,
        "RMSE (Root Mean Squared Error) measures average prediction error — a lower number means more "
        "accurate forecasts. All inflation RMSE values are in percentage points of inflation; "
        "unemployment RMSE values are in percentage points of the unemployment rate."
    )

    # Inflation h=1
    heading(doc, "Inflation — 1 Month Ahead", level=2)
    add_table(doc,
        headers=["Model", "RMSE", "Notes"],
        rows=[
            ["Naive (baseline)",       "0.449", "Simply repeating last month's value"],
            ["ARIMA",                  "0.402", "Classical time-series; modest improvement"],
            ["VAR",                    "0.431", "Multi-variable classical model"],
            ["Ridge Regression",       "0.533", "Hurt by multicollinearity in 84-feature set"],
            ["Lasso Regression ★",    "0.278", "Best overall — automatic variable selection pays off"],
            ["Random Forest",          "0.421", "Solid but not best at 1-month horizon"],
            ["XGBoost",                "0.335", "Strong but Lasso edges it at close range"],
        ],
        col_widths=[1.9, 1.0, 3.4],
        header_bg="2E75B6",
    )
    body(doc,
        "At the shortest horizon, Lasso was the clear winner, beating every other model by a wide margin. "
        "A simple linear model outperforming tree-based methods at close range is somewhat surprising. "
        "The likely explanation: 1-month-ahead inflation is mostly 'wherever it is now plus a small "
        "correction,' and Lasso's variable-selection ability identifies a clean handful of signals without overfitting.",
        italic=True
    )

    # Inflation h=3 and h=6
    heading(doc, "Inflation — 3 and 6 Months Ahead", level=2)
    add_table(doc,
        headers=["Model", "h=3 RMSE", "h=6 RMSE"],
        rows=[
            ["Naive (baseline)",    "0.977", "1.436"],
            ["ARIMA",               "0.957", "1.448"],
            ["VAR",                 "1.127", "1.974"],
            ["Ridge Regression",    "1.424", "1.597"],
            ["Lasso Regression",    "0.934", "1.242"],
            ["Random Forest",       "0.795", "1.060"],
            ["XGBoost ★",          "0.746", "0.831"],
        ],
        col_widths=[2.2, 1.5, 1.5],
        header_bg="2E75B6",
    )
    body(doc,
        "As the horizon extended, the story shifted decisively toward tree-based machine learning. "
        "At 6 months, XGBoost (0.831) was roughly 42% more accurate than ARIMA (1.448). VAR actively "
        "fell apart at longer horizons, performing worse than simply guessing no change.",
        italic=True
    )

    # Unemployment
    heading(doc, "Unemployment Rate — All Horizons", level=2)
    body(doc,
        "Unemployment told an even more dramatic story. In March–April 2020, the unemployment rate "
        "went from 3.5% to 14.7% in two months — the fastest rise in recorded history. No model "
        "trained on the prior 30 years had ever seen anything remotely like this. This single shock "
        "dominated the error statistics for classical models."
    )
    add_table(doc,
        headers=["Model", "h=1 RMSE", "h=3 RMSE", "h=6 RMSE"],
        rows=[
            ["Naive (baseline)",    "0.642",  "1.087",   "1.400"],
            ["ARIMA",               "0.996",  "2.741",   "5.756"],
            ["VAR",                 "0.989",  "2.729",   "5.937"],
            ["Ridge Regression",    "2.211",  "2.338",   "2.665"],
            ["Lasso Regression",    "2.099",  "2.073",   "2.506"],
            ["Random Forest",       "0.448",  "0.903",   "1.168"],
            ["XGBoost ★",          "0.259",  "0.903",   "0.956"],
        ],
        col_widths=[2.2, 1.2, 1.2, 1.2],
        header_bg="2E75B6",
    )
    body(doc,
        "ARIMA's 6-month RMSE of 5.76 means it was, on average, off by nearly 6 percentage points. "
        "At its worst (April–May 2020), it predicted unemployment around 4% when it was actually 14–15%. "
        "XGBoost's error over the same period was around 0.96 percentage points — roughly six times "
        "more accurate. Linear ML models also failed badly: they extrapolate in straight lines and "
        "couldn't pull forecasts back to reality during an unprecedented shock. Tree-based models "
        "naturally cap predictions at values they've seen before, which paradoxically helped them "
        "recover faster from the spike.",
        italic=True
    )

    # Sub-period
    heading(doc, "Sub-Period Findings (Removing COVID)", level=2)
    body(doc,
        "When we looked only at 2008–2019 (a calmer period), the gaps narrowed considerably:"
    )
    bullet(doc, "For unemployment at 1 month ahead, Lasso (0.066) and Ridge (0.123) actually beat "
                "XGBoost (0.243) — linear models work fine in stable regimes.")
    bullet(doc, "For inflation at 3–6 months, tree-based models still led, but by smaller margins.")
    bullet(doc, "The machine learning advantage is not a constant — it depends heavily on the "
                "economic environment.")

    divider(doc)

    # ── Summary ───────────────────────────────────────────────────────────────
    heading(doc, "Summary of Findings")
    findings = [
        ("Machine learning wins overall, but not uniformly.",
         "The advantage is clearest at longer horizons (3–6 months) and during economic shocks."),
        ("XGBoost was the most consistent top performer.",
         "It led across nearly every target and horizon combination."),
        ("Lasso was the surprise winner for 1-month inflation.",
         "Its automatic variable selection gave it an edge in short-run signal extraction."),
        ("Classical models failed catastrophically on unemployment during COVID.",
         "ARIMA produced errors up to 15 percentage points; tree-based models remained far more robust."),
        ("The Naive benchmark is hard to beat for unemployment at 1 month.",
         "Unemployment changes very slowly in calm periods, so predicting no change is genuinely competitive."),
        ("Alpha tuning did not improve results.",
         "The originally chosen regularization values were already near-optimal — itself a finding about "
         "the robustness of those defaults."),
    ]
    for i, (title, desc) in enumerate(findings, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.2)
        r = p.add_run(f"{i}. {title}  ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GRAY

    doc.add_paragraph()
    divider(doc)

    # ── Closing ───────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("The Bottom Line")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY

    body(doc,
        "If you were a policymaker or economist building a real forecasting system today, you would use "
        "XGBoost or Random Forest as your primary models for medium-to-long horizons, supplement with "
        "Lasso for near-term inflation, and keep ARIMA around as a quick sanity-check baseline — while "
        "knowing that any purely statistical model will struggle when history offers no precedent for "
        "what just happened."
    )

    os.makedirs("Debriefs", exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    build()
