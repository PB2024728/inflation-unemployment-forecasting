"""
Generates the revised 10-day project plan as a formatted Word document.
Saved to Debriefs/Inflation_Unemployment_Forecasting_Project_Plan.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join("Debriefs", "Inflation_Unemployment_Forecasting_Project_Plan.docx")

NAVY   = RGBColor(0x1F, 0x39, 0x64)
STEEL  = RGBColor(0x2E, 0x75, 0xB6)
GRAY   = RGBColor(0x59, 0x59, 0x59)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x37, 0x86, 0x44)
AMBER  = RGBColor(0xBF, 0x87, 0x00)
LGRAY  = RGBColor(0x99, 0x99, 0x99)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "1F3964")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        fill = "F2F7FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                text, bold, color = val
                run = p.add_run(text)
                run.bold = bold
                run.font.color.rgb = color
            else:
                run = p.add_run(str(val))
                run.font.color.rgb = GRAY
                if c_idx == 0:
                    run.bold = True
            run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = STEEL
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = STEEL
    return p

def status_badge(text, color_hex):
    """Returns a (text, bold, RGBColor) tuple for use in table cells."""
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    return (text, True, RGBColor(r, g, b))

def divider(doc):
    doc.add_paragraph()

# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title.add_run("Inflation & Unemployment Forecasting Bake-Off")
t.bold = True
t.font.size = Pt(22)
t.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub.add_run("Revised 10-Day Project Plan  |  May 2026")
s.font.size = Pt(11)
s.font.color.rgb = STEEL

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
n = note.add_run("Updated to reflect dashboard output format decision made after Day 6.")
n.italic = True
n.font.size = Pt(9)
n.font.color.rgb = LGRAY

doc.add_paragraph()

# ── Section 1: Project Goal ────────────────────────────────────────────────────
h1(doc, "1. Project Goal")
body(doc,
    "Build a forecasting system that compares traditional time-series models against machine learning "
    "models for predicting inflation and unemployment. The goal is not simply to achieve the lowest "
    "forecast error, but to demonstrate which model class performs better, under what conditions, "
    "and why — and to present those findings in a way that is rigorous, honest, and visually compelling."
)
callout(doc,
    "\"Do traditional time-series models outperform machine learning methods in forecasting inflation "
    "and unemployment, or do ML methods gain an edge once richer predictor sets and nonlinear patterns "
    "are included?\""
)

# ── Section 2: Targets, Horizons, Models ──────────────────────────────────────
h1(doc, "2. Targets, Horizons & Model Lineup")

h2(doc, "Forecast Targets")
add_table(doc,
    ["Variable", "Transformation", "Role"],
    [
        ["CPI Inflation (Headline)", "Year-over-year % change (12-month)", "Primary target"],
        ["Core CPI Inflation",       "Year-over-year % change (12-month)", "Secondary target / predictor"],
        ["Unemployment Rate",        "Level",                              "Primary target"],
    ],
    col_widths=[2.0, 2.3, 1.8]
)
divider(doc)

h2(doc, "Forecast Horizons")
body(doc, "Forecasts are generated at three horizons: 1 month ahead (h=1), 3 months ahead (h=3), and 6 months ahead (h=6).")

h2(doc, "Model Lineup")
add_table(doc,
    ["Category", "Model", "Description"],
    [
        ["Baseline",         "Naive",         "Last observed value carried forward. The minimum bar every model must clear."],
        ["Classical",        "ARIMA",         "Captures autocorrelation in a single series using its own history and forecast errors."],
        ["Classical",        "VAR",           "Models inflation and unemployment jointly, capturing cross-variable dynamics."],
        ["Machine Learning", "Ridge",         "Penalised linear regression across all macro predictors. Controls overfitting via L2 penalty."],
        ["Machine Learning", "Lasso",         "Like Ridge but performs automatic variable selection via L1 penalty."],
        ["Machine Learning", "Random Forest", "Ensemble of 100 decision trees. Robust to outliers and nonlinear relationships."],
        ["Machine Learning", "XGBoost",       "Gradient boosting algorithm. Industry standard for tabular forecasting competitions."],
    ],
    col_widths=[1.5, 1.4, 3.7]
)
divider(doc)

h2(doc, "Evaluation Framework")
bullet(doc, "Expanding-window backtesting from January 2000 through February 2026 (~300 evaluation points per model).", "Method:  ")
bullet(doc, "Minimum 60 months of training data before the first forecast.", "Min training:  ")
bullet(doc, "RMSE (Root Mean Squared Error) and MAE (Mean Absolute Error), reported separately per target and horizon.", "Metrics:  ")
bullet(doc, "No look-ahead bias — each model trained only on data available at that point in time.", "Leakage:  ")

# ── Section 3: Deliverables ────────────────────────────────────────────────────
h1(doc, "3. Final Deliverables")
add_table(doc,
    ["Deliverable", "Format", "Description"],
    [
        ["Interactive Results Dashboard", "Single self-contained .html file",    "Plotly dashboard with model toggles, horizon dropdowns, sub-period analysis, and feature importance"],
        ["Progress Report (Days 1–4)",    "Word document (.docx)",               "Client-facing write-up of scope, data, EDA, and classical model results"],
        ["Final Report",                  "Word document (.docx)",               "Full write-up including ML results, findings, and conclusions"],
        ["Reproducible Codebase",         "Python scripts + .env + README",      "End-to-end pipeline from raw FRED data to final outputs, runnable from one API key"],
        ["Master Error Table",            "CSV (results/master_errors.csv)",     "RMSE and MAE for all 7 models × 2 targets × 3 horizons"],
    ],
    col_widths=[2.2, 2.0, 2.4]
)

# ── Section 4: Day-by-Day Plan ─────────────────────────────────────────────────
h1(doc, "4. Day-by-Day Plan")
body(doc,
    "The plan below reflects the revised structure agreed after Day 6, including the decision to "
    "replace static matplotlib output (originally planned for Day 9) with an interactive Plotly "
    "HTML dashboard."
)

# ── Day 1 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 1 — Project Scoping  ✅  Complete")
body(doc, "Lock in all design decisions before touching any data or code.")
add_table(doc,
    ["Decision", "Choice Made"],
    [
        ["Inflation measure",    "Year-over-year % change in CPI (12-month rolling)"],
        ["Forecast horizons",    "h = 1, 3, 6 months ahead"],
        ["Forecast frequency",   "Monthly"],
        ["Evaluation method",    "Expanding window from January 2000"],
        ["Min training window",  "60 months (5 years)"],
    ],
    col_widths=[2.3, 4.3]
)
divider(doc)
bullet(doc, "One-paragraph project objective", "Deliverable:  ")
bullet(doc, "Locked list of target variables, forecast horizons, and initial model list")

# ── Day 2 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 2 — Data Collection  ✅  Complete")
body(doc,
    "Pull all macroeconomic series from FRED, align to monthly frequency, and apply transformations. "
    "Tools: fredapi, pandas, python-dotenv."
)
add_table(doc,
    ["Series", "FRED ID", "Transformation"],
    [
        ["CPI (All Items)",       "CPIAUCSL",   "YoY % change"],
        ["Core CPI",              "CPILFESL",   "YoY % change"],
        ["Unemployment Rate",     "UNRATE",     "Level"],
        ["Federal Funds Rate",    "FEDFUNDS",   "Level"],
        ["10-Year Treasury",      "GS10",       "Level"],
        ["3-Month T-Bill",        "TB3MS",      "Level"],
        ["Yield Spread (10Y−3M)", "Derived",    "Level"],
        ["Industrial Production", "INDPRO",     "MoM % change"],
        ["Nonfarm Payrolls",      "PAYEMS",     "MoM % change"],
        ["Oil Prices (WTI)",      "DCOILWTICO", "MoM % change"],
        ["Consumer Sentiment",    "UMCSENT",    "Level"],
        ["M2 Money Supply",       "M2SL",       "MoM % change"],
    ],
    col_widths=[2.0, 1.5, 3.1]
)
divider(doc)
bullet(doc, "data/raw/raw_macro.csv — 435 rows, 11 series, 1990–2026", "Output:  ")
bullet(doc, "data/processed/macro_monthly.csv — 421 rows, 12 model-ready features")
bullet(doc, "data/data_dictionary.md — full variable definitions")

# ── Day 3 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 3 — Exploratory Data Analysis  ✅  Complete")
body(doc,
    "Understand the structure of the series before modelling. "
    "Tools: matplotlib, seaborn."
)
bullet(doc, "Inflation averaged 2.6% (range: −2.0% to +9.0%). 2022 spike is a clear structural outlier.")
bullet(doc, "Unemployment averaged 5.7% (range: 3.4% to 14.8%). COVID-19 spike in April 2020 is the dominant event.")
bullet(doc, "Phillips Curve relationship has broken down post-2000. 2022 shows high inflation with moderate unemployment simultaneously.")
bullet(doc, "Yield spread (10Y−3M) inverted ahead of both 2008 and 2020 recessions — a valuable leading indicator.")
bullet(doc, "Oil prices show extreme monthly volatility (−54% to +85%), validating the percentage-change transformation.")
divider(doc)
bullet(doc, "Six publication-quality figures saved to figures/", "Deliverable:  ")

# ── Day 4 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 4 — Classical Benchmark Models  ✅  Complete")
body(doc,
    "Build and backtest Naive, ARIMA, and VAR using the expanding-window framework. "
    "Tools: statsmodels."
)
add_table(doc,
    ["Model", "Inflation h=1", "Inflation h=3", "Inflation h=6", "Unemp. h=1", "Unemp. h=3", "Unemp. h=6"],
    [
        ["Naive", "0.449", "0.977", "1.435", "0.642", "1.087", "1.400"],
        ["ARIMA", "0.402", "0.979", "1.460", "1.039", "2.951", "7.003"],
        ["VAR",   "0.431", "1.127", "1.974", "0.989", "2.729", "5.937"],
    ],
    col_widths=[1.0, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95]
)
divider(doc)
body(doc,
    "Key finding: ARIMA edges Naive at h=1 for inflation (+10%) but is beaten by Naive at longer horizons. "
    "For unemployment, Naive dominates at every horizon. The ARIMA h=6 unemployment RMSE of 7.0 is driven "
    "by the COVID-19 shock — ARIMA extrapolated the spike as a persistent trend, producing catastrophic errors."
)
bullet(doc, "results/forecasts_arima_var.csv and results/errors_arima_var.csv", "Deliverable:  ")

# ── Day 5 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 5 — Machine Learning Models  ✅  Complete")
body(doc,
    "Build and backtest Ridge, Lasso, Random Forest, and XGBoost with 84 lagged macro features. "
    "Tools: scikit-learn, xgboost."
)
add_table(doc,
    ["Model", "Inflation h=1", "Inflation h=3", "Inflation h=6", "Unemp. h=1", "Unemp. h=3", "Unemp. h=6"],
    [
        ["Ridge",  "0.514", "1.359", "1.545", "2.150", "2.261", "2.615"],
        ["Lasso",  "0.229", "0.811", "1.163", "1.902", "2.057", "2.564"],
        ["RF",     "0.421", "0.795", "1.060", "0.448", "0.903", "1.168"],
        ["XGBoost","0.356", "0.769", "0.859", "0.296", "0.900", "0.956"],
    ],
    col_widths=[1.0, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95]
)
divider(doc)
body(doc,
    "Key findings: Lasso dominates at h=1 inflation (0.229 vs ARIMA 0.402). XGBoost wins at all other "
    "cells. Tree-based models handle the COVID spike far better than linear ML — Ridge and Lasso are "
    "worse than Naive for unemployment at every horizon."
)
bullet(doc, "results/forecasts_ml.csv and results/errors_ml.csv", "Deliverable:  ")

# ── Day 6 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 6 — Master Bake-Off Merge  ✅  Complete")
body(doc,
    "Merge all forecast outputs into a single master table for comparison across all 7 models, "
    "2 targets, and 3 horizons."
)
add_table(doc,
    ["", "h=1 Winner", "h=3 Winner", "h=6 Winner"],
    [
        ["Inflation",    "Lasso (0.229)",    "XGBoost (0.769)", "XGBoost (0.859)"],
        ["Unemployment", "XGBoost (0.296)",  "XGBoost (0.900)", "XGBoost (0.956)"],
    ],
    col_widths=[1.3, 1.8, 1.8, 1.8]
)
divider(doc)
bullet(doc, "results/master_forecasts.csv — 1,864 forecast observations", "Deliverable:  ")
bullet(doc, "results/master_errors.csv — RMSE and MAE for all 42 model-target-horizon combinations")

# ── Day 7 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 7 — Diagnostics & Interpretation  ⬜  Upcoming")
body(doc,
    "Move beyond 'which model won' and understand why. "
    "Outputs from this day feed directly into the dashboard panels."
)
bullet(doc, "Residual plots per model — check for systematic patterns or structural break sensitivity")
bullet(doc, "Feature importance for Random Forest and XGBoost — which predictors drive each target")
bullet(doc, "Sub-period RMSE breakdown — pre-2008, 2008–2019, 2020+ — to assess stability across regimes")
bullet(doc, "All outputs saved as CSVs for dashboard consumption")
divider(doc)
bullet(doc, "results/subperiod_errors.csv", "Deliverable:  ")
bullet(doc, "results/feature_importance_rf.csv and results/feature_importance_xgb.csv")
bullet(doc, "Residual analysis figures (figures/residuals_*.png)")

# ── Day 8 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 8 — Refinement  ⬜  Upcoming")
body(doc,
    "Review the diagnostic outputs and tighten the weakest components. Focus on improving the quality "
    "of the comparison, not adding new models."
)
bullet(doc, "Tune ARIMA order selection if residuals show systematic patterns")
bullet(doc, "Tune XGBoost and RF hyperparameters using a held-out validation window")
bullet(doc, "Review Lasso penalty — confirm it is not over-shrinking at longer horizons")
bullet(doc, "Assess whether any predictors should be dropped based on feature importance findings")
divider(doc)
bullet(doc, "Final confirmed model specifications for all 7 models", "Deliverable:  ")

# ── Day 9 ──────────────────────────────────────────────────────────────────────
h2(doc, "Day 9 — Interactive Dashboard  ⬜  Upcoming  [Revised from original plan]")
body(doc,
    "Build a single self-contained interactive HTML dashboard using Plotly. "
    "This replaces the static matplotlib visualisation script originally planned for Day 9. "
    "Tools: plotly."
)
add_table(doc,
    ["Dashboard Panel", "Chart Type", "Interactive Controls"],
    [
        ["Winner scorecard",       "Summary table",          "None — fixed overview"],
        ["RMSE horse race",        "Bar chart",              "Target dropdown, horizon dropdown"],
        ["Actual vs Predicted",    "Time series",            "Model toggles, horizon selector, zoom"],
        ["Sub-period breakdown",   "Grouped bar chart",      "Target dropdown, horizon dropdown"],
        ["Feature importance",     "Horizontal bar chart",   "Model selector (RF vs XGBoost), target + horizon dropdowns"],
    ],
    col_widths=[1.9, 1.6, 3.1]
)
divider(doc)
bullet(doc, "dashboard.html — single file, no server required, fully portable", "Deliverable:  ")

# ── Day 10 ─────────────────────────────────────────────────────────────────────
h2(doc, "Day 10 — Final Report  ⬜  Upcoming")
body(doc,
    "Produce the final written report, update the progress Word document with complete findings, "
    "and finalise the project README."
)
body(doc, "Report structure:")
bullet(doc, "Motivation — why forecasting inflation and unemployment matters")
bullet(doc, "Research question — traditional vs ML methods")
bullet(doc, "Data — sources, transformations, sample period")
bullet(doc, "Methods — all 7 models with brief technical descriptions")
bullet(doc, "Evaluation — expanding-window framework, RMSE, MAE")
bullet(doc, "Findings — which models won, at which horizons, and why")
bullet(doc, "Conclusion — does ML add value, or do simpler models remain competitive?")
divider(doc)
bullet(doc, "Final Word report (.docx) with complete results and conclusions", "Deliverable:  ")
bullet(doc, "Updated README.md with final findings summary")

# ── Section 5: Technology Stack ────────────────────────────────────────────────
h1(doc, "5. Technology Stack")
add_table(doc,
    ["Library", "Version", "Purpose"],
    [
        ["pandas",         "3.0.2",  "Data wrangling and time series alignment"],
        ["numpy",          "2.4.4",  "Numerical operations and error metrics"],
        ["fredapi",        "0.5.2",  "FRED API data pull"],
        ["statsmodels",    "0.14.6", "ARIMA and VAR models"],
        ["scikit-learn",   "1.8.0",  "Ridge, Lasso, Random Forest, pipeline utilities"],
        ["xgboost",        "3.2.0",  "XGBoost gradient boosting"],
        ["matplotlib",     "3.10.9", "EDA figures (Days 2–3)"],
        ["seaborn",        "0.13.2", "Correlation heatmaps and styled EDA plots"],
        ["plotly",         "TBD",    "Interactive HTML dashboard (Day 9)"],
        ["python-docx",    "1.2.0",  "Word document generation"],
        ["python-dotenv",  "1.2.2",  "Secure API key management"],
    ],
    col_widths=[1.6, 0.9, 4.1]
)

# ── Section 6: Repo Structure ──────────────────────────────────────────────────
h1(doc, "6. Repository Structure")
add_table(doc,
    ["Path", "Contents"],
    [
        ["data/raw/",                    "Raw FRED CSV downloads"],
        ["data/processed/",              "Clean merged monthly feature dataset"],
        ["data/data_dictionary.md",      "Variable definitions and transformations"],
        ["src/01_collect_data.py",       "Day 2: FRED pull and feature engineering"],
        ["src/02_eda.py",                "Day 3: Six EDA figures"],
        ["src/03_arima_var.py",          "Day 4: ARIMA + VAR expanding-window backtest"],
        ["src/04_ml_models.py",          "Day 5: Ridge / Lasso / RF / XGBoost backtest"],
        ["src/05_backtesting.py",        "Day 6: Master error table merge"],
        ["src/06_diagnostics.py",        "Day 7: Residuals, feature importance, sub-period"],
        ["src/07_dashboard.py",          "Day 9: Plotly interactive HTML dashboard"],
        ["results/",                     "Forecast CSVs and error tables"],
        ["figures/",                     "EDA and diagnostic plots"],
        ["Debriefs/",                    "Word documents — progress reports and project plan"],
        [".env",                         "FRED API key (not committed to version control)"],
        ["requirements.txt",             "Full dependency list"],
        ["README.md",                    "Project overview and run instructions"],
    ],
    col_widths=[2.8, 3.8]
)

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    "All code, data, and outputs are reproducible from a single FRED API key with no manual data entry."
)
footer_run.italic = True
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = LGRAY

doc.save(OUT)
print(f"Document saved: {OUT}")
